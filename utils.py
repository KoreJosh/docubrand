import os
from PyPDF2 import PdfFileReader, PdfFileWriter
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from io import BytesIO
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

def overlay_images_on_pdf_page(page, logo_path=None, footer_path=None, page_size=letter):
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=page_size)

    if logo_path and os.path.exists(logo_path):
        logo = Image.open(logo_path)
        logo_width = 100
        logo_height = int(logo.height * (logo_width / logo.width))
        can.drawImage(logo_path, 40, page_size[1] - 50, width=logo_width, height=logo_height)

    if footer_path and os.path.exists(footer_path):
        footer = Image.open(footer_path)
        footer_width = 120
        footer_height = int(footer.height * (footer_width / footer.width))
        can.drawImage(footer_path, 40, 20, width=footer_width, height=footer_height)

    can.save()
    packet.seek(0)
    new_pdf = PdfFileReader(packet)
    overlay_page = new_pdf.getPage(0)
    page.merge_page(overlay_page)
    return page

def process_pdf(input_path, logo_path, footer_path, output_path, split_pages=False):
    pdf_reader = PdfFileReader(open(input_path, "rb"))
    pdf_writer = PdfFileWriter()

    for page_num in range(pdf_reader.getNumPages()):
        page = pdf_reader.getPage(page_num)
        branded_page = overlay_images_on_pdf_page(page, logo_path, footer_path)
        pdf_writer.addPage(branded_page)

    with open(output_path, "wb") as out_file:
        pdf_writer.write(out_file)

    if split_pages:
        os.makedirs("split_pages", exist_ok=True)
        for i in range(pdf_reader.getNumPages()):
            writer = PdfFileWriter()
            writer.addPage(pdf_reader.getPage(i))
            with open(f"split_pages/page_{i+1}.pdf", "wb") as f:
                writer.write(f)

def process_docx(input_path, logo_path, footer_path, output_path, split_pages=False):
    doc = Document(input_path)

    # Add header
    if logo_path and os.path.exists(logo_path):
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.5))

    # Add footer
    if footer_path and os.path.exists(footer_path):
        footer = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(footer_path, width=Inches(1.5))

    docx_path = output_path.replace(".pdf", ".docx")
    doc.save(docx_path)

    convert_docx_to_pdf(docx_path, output_path)

    if split_pages:
        process_pdf(output_path, None, None, output_path, split_pages=True)

    return docx_path, output_path

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    except:
        doc = Document(docx_path)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
        pdf.output(pdf_path)

